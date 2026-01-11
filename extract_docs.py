import os
import glob
from docx import Document
import pandas as pd

def extract_docx(file_path):
    doc = Document(file_path)
    content = [f"# {os.path.basename(file_path)}\n"]
    for para in doc.paragraphs:
        if para.text.strip():
            content.append(para.text)
    
    # Extract tables
    for table in doc.tables:
        content.append("\n## Table")
        rows = []
        for row in table.rows:
            row_data = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
            rows.append("| " + " | ".join(row_data) + " |")
        
        if rows:
            header = rows[0]
            separator = "| " + " | ".join(["---"] * len(rows[0].split("|"))).replace("|", "", 2).strip() + " |" # Approximate logic
            # Simpler table handling
            content.append("\n| " + " | ".join(["Col " + str(i+1) for i in range(len(table.columns))]) + " |")
            content.append("| " + " | ".join(["---"] * len(table.columns)) + " |")
            content.extend(rows)
            
    return "\n\n".join(content)

def extract_xlsx(file_path):
    xl = pd.ExcelFile(file_path)
    content = [f"# {os.path.basename(file_path)}\n"]
    
    for sheet_name in xl.sheet_names:
        content.append(f"## Sheet: {sheet_name}\n")
        df = xl.parse(sheet_name)
        content.append(df.to_markdown(index=False))
        content.append("\n---\n")
        
    return "\n".join(content)

def main():
    base_dir = "Reference-Docs"
    output_dir = "Reference-Docs/Extracted"
    os.makedirs(output_dir, exist_ok=True)
    
    # Process DOCX
    for docx_file in glob.glob(os.path.join(base_dir, "*.docx")):
        try:
            print(f"Processing {docx_file}...")
            text = extract_docx(docx_file)
            out_name = os.path.basename(docx_file) + ".md"
            with open(os.path.join(output_dir, out_name), "w", encoding="utf-8") as f:
                f.write(text)
        except Exception as e:
            print(f"Error processing {docx_file}: {e}")

    # Process XLSX
    for xlsx_file in glob.glob(os.path.join(base_dir, "*.xlsx")):
        try:
            print(f"Processing {xlsx_file}...")
            text = extract_xlsx(xlsx_file)
            out_name = os.path.basename(xlsx_file) + ".md"
            with open(os.path.join(output_dir, out_name), "w", encoding="utf-8") as f:
                f.write(text)
        except Exception as e:
            print(f"Error processing {xlsx_file}: {e}")

if __name__ == "__main__":
    main()
